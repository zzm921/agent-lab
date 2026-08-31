"""全链路端到端测试：混合好坏文档 → 报告/DLQ/干净文本；建库冒烟（memory 后端）。

二进制文档（PDF/DOCX/GBK/加密 PDF）在测试内现场生成，不向仓库提交二进制夹具。
"""
import io
import json
import os
import shutil
from pathlib import Path

import fitz
import pytest

from app.llm.fake_model import FakeEmbeddings
from app.rag.manager import RagManager
from app.rag.preprocess import run_pipeline
from tests.conftest import make_settings

FIXTURES = Path(__file__).resolve().parents[1] / "tests" / "fixtures" / "preprocess"


def _text_pdf(pages: int = 3) -> bytes:
    """3 页文本 PDF：每页重复页眉 + 正文 + 页码（insert_text 不支持换行，分行插入）。"""
    doc = fitz.open()
    for i in range(pages):
        page = doc.new_page()
        page.insert_text((72, 50), "云帆科技 内部机密 文件编号：ZD-2024-001", fontname="china-s", fontsize=10)
        body = f"第{i + 1}章 差旅制度正文：住宿一线城市每晚不超过四百元，市内交通单日上限八十元，须凭票据实报销。"
        for j in range(0, len(body), 20):
            page.insert_text((72, 100 + (j // 20) * 16), body[j : j + 20], fontname="china-s", fontsize=12)
        page.insert_text((72, 780), f"第 {i + 1} 页 / 共 {pages} 页", fontname="china-s", fontsize=10)
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("会议室管理规定", level=1)
    doc.add_paragraph("会议室使用须提前一天在 OA 系统预约，会后须清理白板并关闭设备。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "会议室"
    table.rows[0].cells[1].text = "容量"
    table.rows[1].cells[0].text = "A101"
    table.rows[1].cells[1].text = "12 人"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _scan_pdf() -> bytes:
    doc = fitz.open()
    doc.new_page()  # 空白页（无文本层 → 路由 OCR）
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


@pytest.fixture
def input_dir(tmp_path, monkeypatch):
    """构造混合输入目录：好文档/重复/乱码/低质/伪装/扫描件/加密。"""
    src = tmp_path / "docs"
    src.mkdir()
    # 文本夹具从仓库夹具目录复制
    for name in ("good_doc.md", "dup_v1.md", "dup_v2.md", "garbled.txt", "log_fragment.txt", "fake_pdf.txt"):
        shutil.copy(FIXTURES / name, src / name)
    # 二进制夹具现场生成
    (src / "meeting.docx").write_bytes(_docx())
    (src / "policy.pdf").write_bytes(_text_pdf())
    (src / "scan.pdf").write_bytes(_scan_pdf())
    (src / "gbk_legacy.txt").write_bytes(
        "这是一份历史遗留的 GBK 编码文档，用于验证编码检测能力。"
        "员工考勤以系统记录为准，迟到早退均需在当月内完成核对处理。"
        "年假按工龄计算，工龄满一年不满十年的享有五天年假，满十年的享有十天年假。"
        "加班须事先申请审批，未经审批的加班时间不予结算加班费。"
        "本制度自发布之日起施行，由人力资源部负责解释与修订。".encode("gbk")
    )
    # 控制 mtime：dup_v1 旧、dup_v2 新（去重保留新版本）
    os.utime(src / "dup_v1.md", (1000, 1000))
    os.utime(src / "dup_v2.md", (2000, 2000))
    return src


@pytest.fixture
def ocr_mock(monkeypatch):
    # 足够长的多句 OCR 文本，保证质量分 ≥70（体量分按 500 字封顶线性计算）
    text = (
        "员工手册总则规定，全员应当遵守考勤与差旅各项制度安排。"
        "员工因公出差需提前三天在 OA 系统提交差旅申请，注明出差地点、事由与预计费用。"
        "住宿标准按职级执行，普通员工一线城市每晚不超过四百元，其他城市不超过三百元。"
        "市内交通费凭票据实报销，单日上限八十元，出差期间原则上不安排宴请活动。"
        "报销材料须在返回后十个工作日内提交，包括行程单、发票与审批单，逾期视为放弃。"
    )
    monkeypatch.setattr(
        "app.rag.preprocess.parsers.ocr_parser.ocr_image",
        lambda data: text,
    )


class TestRunPipeline:
    def test_statuses_and_outputs(self, input_dir, tmp_path, ocr_mock):
        out = tmp_path / "out"
        clean_docs, report = run_pipeline(input_dir, out)
        summary = report.summary()

        by_name = {Path(d.path).name: d for d in report.docs}
        assert by_name["good_doc.md"].status == "ok"
        assert by_name["dup_v2.md"].status == "ok"
        assert by_name["dup_v1.md"].status == "superseded"  # 旧版本被覆盖
        assert by_name["garbled.txt"].status == "dlq"  # 乱码拦截
        assert by_name["fake_pdf.txt"].status == "dlq"  # 伪装格式损坏 PDF
        assert by_name["log_fragment.txt"].status == "dlq"  # 质量分过低
        assert by_name["policy.pdf"].status == "ok"
        assert by_name["scan.pdf"].status == "ok"  # OCR mock 后可入库
        assert summary["ok"] >= 4 and summary["dlq"] >= 3

        # 入库列表不含被去重/淘汰文档
        sources = [Path(cd.metadata["source"]).name for cd in clean_docs]
        assert "dup_v1.md" not in sources and "good_doc.md" in sources

        # 页眉页脚已移除：PDF 清洗文本不含「内部机密」与页码
        pdf_doc = next(cd for cd in clean_docs if cd.metadata["source"].endswith("policy.pdf"))
        assert "内部机密" not in pdf_doc.text and "共 3 页" not in pdf_doc.text

        # GBK 文档被正确解码（非乱码拦截）
        assert by_name["gbk_legacy.txt"].status == "ok"

        # 报告落盘 + DLQ 归档
        report_data = json.loads((out / "report.json").read_text(encoding="utf-8"))
        assert report_data["summary"]["total"] == len(report.docs)
        dlq_names = {p.name for p in (out / "dlq").iterdir()}
        assert "garbled.txt" in dlq_names and "garbled.txt.error.txt" in dlq_names

    def test_empty_input_dir(self, tmp_path):
        out = tmp_path / "out"
        clean_docs, report = run_pipeline(tmp_path / "docs", out)
        assert clean_docs == [] and report.summary()["total"] == 0


class TestIngestSmoke:
    def test_clean_docs_ingest_to_memory_store(self, input_dir, tmp_path, ocr_mock):
        clean_docs, _ = run_pipeline(input_dir, tmp_path / "out")
        settings = make_settings(rag_schemes=["naive"])  # 未配 Qdrant/ES → 内存后端
        rag = RagManager(settings, FakeEmbeddings(), top_k=3)
        rag.ingest_all([cd.text for cd in clean_docs])
        assert len(rag.get("naive")) == len(clean_docs)
