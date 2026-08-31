"""解析器测试：md/html 结构、docx 标题与表格、文本 PDF 页序、OCR 路由（mock 不联网）。"""
import io
from pathlib import Path

import fitz
import pytest

from app.rag.preprocess.models import RawFile
from app.rag.preprocess.parsers import route_parse
from app.rag.preprocess.parsers.docx_parser import parse_docx
from app.rag.preprocess.parsers.ocr_parser import parse_ocr
from app.rag.preprocess.parsers.pdf_parser import parse_pdf
from app.rag.preprocess.parsers.text_parser import parse_text
from app.rag.preprocess.sniffer import sniff


def _raw(data: bytes, name: str) -> RawFile:
    return sniff(data, Path(name))


class TestTextParser:
    def test_markdown_structure(self):
        md = "# 第一章 总则\n\n员工应当遵守制度。\n\n| 职级 | 限额 |\n| 普通员工 | 400 |\n\n## 第二节 报销\n\n凭票据实报销。".encode()
        parsed = parse_text(_raw(md, "a.md"))
        types = [(el.type, el.level) for el in parsed.elements]
        assert ("title", 1) in types and ("title", 2) in types
        assert any(el.type == "table" and "职级" in el.text for el in parsed.elements)

    def test_html_tags_stripped(self):
        html = b"<html><body><h1>\xe6\xa0\x87\xe9\xa2\x98</h1><p>\xe6\xad\xa3\xe6\x96\x87\xe5\x86\x85\xe5\xae\xb9</p></body></html>"
        parsed = parse_text(_raw(html, "a.html"))
        text = "\n".join(el.text for el in parsed.elements)
        assert "<p>" not in text and "正文内容" in text


class TestDocxParser:
    def test_headings_table_in_order(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_heading("第一章 总则", level=1)
        doc.add_paragraph("员工应当遵守考勤规定，按时上下班打卡。")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "职级"
        table.rows[0].cells[1].text = "住宿限额"
        table.rows[1].cells[0].text = "普通员工"
        table.rows[1].cells[1].text = "400 元"
        doc.add_heading("第二节 报销", level=2)
        doc.add_paragraph("凭票据实报销。")
        path = tmp_path / "t.docx"
        doc.save(path)

        parsed = parse_docx(_raw(path.read_bytes(), "t.docx"))
        kinds = [el.type for el in parsed.elements]
        assert kinds == ["title", "text", "table", "title", "text"]
        title = parsed.elements[0]
        assert title.level == 1
        table_text = parsed.elements[2].text
        assert "职级: 普通员工" in table_text and "住宿限额: 400 元" in table_text


class TestPdfParser:
    def test_pages_and_markers(self, tmp_path):
        doc = fitz.open()
        for i in range(2):
            page = doc.new_page()
            page.insert_text((72, 72), f"第{i + 1}页正文：制度内容若干。", fontname="china-s")
        import io

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        doc.close()

        parsed = parse_pdf(_raw(data, "t.pdf"))
        assert parsed.page_count == 2
        markers = [el for el in parsed.elements if el.type == "page_marker"]
        assert len(markers) == 2
        assert any("第1页正文" in el.text for el in parsed.elements)


class TestOcrParser:
    def test_image_ocr_mocked(self, monkeypatch):
        monkeypatch.setattr(
            "app.rag.preprocess.parsers.ocr_parser.ocr_image",
            lambda data: "扫描件文字：员工手册总则。\n",
        )
        parsed = parse_ocr(_raw(b"\xff\xd8\xff\xe0fakejpeg", "scan.jpg"))
        assert parsed.route == "ocr"
        assert "员工手册总则" in parsed.elements[0].text

    def test_scanned_pdf_pages_mocked(self, tmp_path, monkeypatch):
        doc = fitz.open()
        doc.new_page()  # 空白页（无文本层）
        doc.new_page()
        import io

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        doc.close()

        calls = []

        def fake_ocr(image_bytes: bytes) -> str:
            calls.append(image_bytes[:4])
            return f"OCR 第 {len(calls)} 页文本。"

        monkeypatch.setattr("app.rag.preprocess.parsers.ocr_parser.ocr_image", fake_ocr)
        parsed = parse_ocr(_raw(data, "scan.pdf"))
        assert len(calls) == 2 and calls[0] == b"\x89PNG"  # 每页渲染为 PNG
        assert parsed.page_count == 2
        assert any("OCR 第 2 页" in el.text for el in parsed.elements)


class TestRouteParse:
    def test_unknown_mime_raises(self):
        # 完整 PNG magic → image/png，light 路由不支持图片 → 报错
        with pytest.raises(ValueError):
            route_parse(_raw(b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR", "x.png"), "light")
